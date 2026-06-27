#!/usr/bin/env python3
"""
gerar_docx.py — Preenche o Template-Offerbook.docx no lugar usando o conteúdo do MD.

Estrategia:
1. Abre o template oficial (mantém capa, índice, formatação).
2. Mapeia cada Heading 1/2/3 do template.
3. Casa cada heading com a seção correspondente no MD do offerbook.
4. Apaga os parágrafos "prompt" (instruções genéricas tipo "Qual o principal benefício?")
   que estão abaixo de cada Heading 3.
5. Insere o conteúdo do MD no lugar.
6. Salva como offerbook-{slug}.docx.

NÃO anexa nada no final. NÃO modifica o template original.

Uso:
    python gerar_docx.py offerbook-cohort-marketing.md
    python gerar_docx.py offerbook-cohort-marketing.md --output meu-offerbook.docx
    python gerar_docx.py offerbook-cohort-marketing.md --template /path/Template-Offerbook.docx

Dependência:
    pip install python-docx
"""

import argparse
import re
import shutil
import sys
import unicodedata
from pathlib import Path


def carregar_python_docx():
    """Importa python-docx ou orienta a instalar."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        return Document, qn
    except ImportError:
        print("=" * 64)
        print("ERRO: python-docx nao esta instalado.")
        print("=" * 64)
        print("\nPara instalar (executar uma vez):\n")
        print("    pip install python-docx\n")
        sys.exit(1)


def encontrar_template(template_arg=None):
    """Busca o Template-Offerbook.docx em locais conhecidos."""
    if template_arg:
        p = Path(template_arg).expanduser().resolve()
        if p.exists():
            return p
        print(f"ERRO: template informado nao existe: {p}")
        sys.exit(1)

    script_dir = Path(__file__).resolve().parent
    candidatos = [
        script_dir.parent / "templates" / "Template-Offerbook.docx",
        script_dir / "Template-Offerbook.docx",
        # Caminho relativo ao cwd (caso aluno rode do projeto)
        Path.cwd() / "aula-01" / "templates" / "Template-Offerbook.docx",
        Path.cwd() / "templates" / "Template-Offerbook.docx",
        # Canonical (skill global)
        Path.home() / ".claude/skills/offerbook/templates/Template-Offerbook.docx",
    ]
    for c in candidatos:
        if c.exists():
            return c
    print("ERRO: Template-Offerbook.docx nao encontrado nos caminhos:")
    for c in candidatos:
        print(f"  - {c}")
    print("\nPasse o caminho com --template /seu/caminho/Template-Offerbook.docx")
    sys.exit(1)


def normalizar(texto):
    """Normaliza string pra comparacao: minuscula, sem acento, sem espaco extra."""
    if not texto:
        return ""
    nfkd = unicodedata.normalize("NFKD", texto)
    sem_acento = "".join(c for c in nfkd if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", sem_acento.lower().strip())


def parsear_md(md_path: Path) -> dict:
    """
    Le o MD do offerbook e devolve dict aninhado:
    {
      'titulo': 'Offerbook — Nome',
      'h1': {
        'Materiais': {
          'h2': {
            'Links Social Media / URLs': {'conteudo': '...', 'h3': {}}
          }
        },
        'Posicionamento': {
          'h2': {
            'Oferta': {
              'h3': {
                'Nome do produto': 'Texto preenchido aqui',
                'Promessa': '...'
              }
            }
          }
        }
      }
    }
    """
    if not md_path.exists():
        print(f"ERRO: arquivo MD nao encontrado: {md_path}")
        sys.exit(1)
    raw = md_path.read_text(encoding="utf-8")

    dados = {"titulo": "", "h1": {}}

    # Titulo (primeiro H1)
    m = re.search(r"^#\s+(.+?)$", raw, re.MULTILINE)
    if m:
        dados["titulo"] = m.group(1).strip()

    # Pega tudo depois do primeiro H1
    if m:
        raw = raw[m.end():]

    # Divide por H2 (##)
    partes_h2 = re.split(r"^##\s+(.+?)$", raw, flags=re.MULTILINE)
    # partes_h2[0] = lixo antes do primeiro H2
    # partes_h2[1] = nome do H2, partes_h2[2] = conteudo, ...

    h1_atual = None
    h2_atual_dict = None

    for i in range(1, len(partes_h2), 2):
        h2_nome = partes_h2[i].strip()
        h2_conteudo = partes_h2[i + 1] if i + 1 < len(partes_h2) else ""

        # H2 do offerbook MD = "Bloco N - Materiais" ou similar.
        # Tenta extrair o nome do bloco "depois do tracinho/em-dash"
        m_bloco = re.match(r"^(?:bloco\s+\d+\s*[-–—]\s*)?(.+)$", h2_nome, re.IGNORECASE)
        nome_limpo = m_bloco.group(1).strip() if m_bloco else h2_nome

        # Determina se esse H2 do MD vira H1 no template (Bloco) ou H2 (subsecao)
        # Bloco do template (H1): Materiais, Posicionamento, Analise de Cliente, Copy
        # Mapeamento por normalizacao
        nome_norm = normalizar(nome_limpo)
        mapeamento_blocos = {
            "materiais": "Materiais",
            "posicionamento": "Posicionamento",
            "oferta": "Posicionamento",  # bloco do MD "Posicionamento > Oferta"
            "posicionamento oferta": "Posicionamento",
            "posicionamento > oferta": "Posicionamento",
            "avatar": "Posicionamento",  # avatar entra em Posicionamento no template
            "historia": "Posicionamento",  # historia entra em Posicionamento
            "historia v2": "Posicionamento",
            "analise de cliente": "Analise de cliente",
            "analise do cliente": "Analise de cliente",
            "copy": "Copy",
        }

        # Identifica H1 (bloco) ao qual esse H2 pertence
        bloco_h1 = mapeamento_blocos.get(nome_norm, None)

        # Se nao mapeou, trata como sub-bloco do bloco atual (ou cria como bloco solto)
        if bloco_h1 is None:
            # Heuristica: se tem palavra-chave conhecida no nome, encaixa
            for chave, valor in mapeamento_blocos.items():
                if chave in nome_norm:
                    bloco_h1 = valor
                    break
            if bloco_h1 is None:
                bloco_h1 = nome_limpo  # ultimo fallback

        if bloco_h1 not in dados["h1"]:
            dados["h1"][bloco_h1] = {"h2": {}}

        # Divide o conteudo do H2 por H3 (###)
        partes_h3 = re.split(r"^###\s+(.+?)$", h2_conteudo, flags=re.MULTILINE)
        # partes_h3[0] = conteudo do H2 antes de qualquer H3 (vai como descricao do H2)

        h2_descricao = partes_h3[0].strip() if partes_h3 else ""

        # Nome do H2 no template (usa o nome limpo do MD)
        h2_template_nome = nome_limpo

        if h2_template_nome not in dados["h1"][bloco_h1]["h2"]:
            dados["h1"][bloco_h1]["h2"][h2_template_nome] = {
                "descricao": h2_descricao,
                "h3": {},
            }

        for j in range(1, len(partes_h3), 2):
            h3_nome = partes_h3[j].strip()
            h3_conteudo = partes_h3[j + 1].strip() if j + 1 < len(partes_h3) else ""
            dados["h1"][bloco_h1]["h2"][h2_template_nome]["h3"][h3_nome] = h3_conteudo

    return dados


def buscar_conteudo(dados: dict, nome_normalizado: str, nivel: str = "h3"):
    """
    Procura conteudo no dict de dados por nome normalizado.
    Procura em qualquer nivel (h1/h2/h3) e devolve string com conteudo, ou None.
    """
    for h1_nome, h1_dados in dados.get("h1", {}).items():
        if nivel == "h1" and normalizar(h1_nome) == nome_normalizado:
            return None  # H1 nao tem conteudo direto

        for h2_nome, h2_dados in h1_dados.get("h2", {}).items():
            if nivel == "h2" and normalizar(h2_nome) == nome_normalizado:
                return h2_dados.get("descricao", "")

            for h3_nome, h3_conteudo in h2_dados.get("h3", {}).items():
                if nivel == "h3" and normalizar(h3_nome) == nome_normalizado:
                    return h3_conteudo

                # Fallback: procura por substring
                if nivel == "h3" and (nome_normalizado in normalizar(h3_nome) or normalizar(h3_nome) in nome_normalizado):
                    return h3_conteudo
    return None


def deletar_paragrafo(p):
    """Remove um paragrafo do doc preservando o XML pai."""
    elem = p._element
    elem.getparent().remove(elem)


def inserir_conteudo_apos(paragrafo_anchor, conteudo_texto, doc):
    """
    Insere conteudo (multiplas linhas) como novos paragrafos logo APOS o paragrafo_anchor.
    Preserva ordem: cada linha do conteudo vira um paragrafo Normal.
    Usa o XML pra inserir na posicao certa.
    """
    if not conteudo_texto or not conteudo_texto.strip():
        return

    from docx.oxml.ns import qn
    from copy import deepcopy

    linhas = [l.strip() for l in conteudo_texto.split("\n") if l.strip()]

    elem_anchor = paragrafo_anchor._element
    parent = elem_anchor.getparent()
    indice_anchor = list(parent).index(elem_anchor)

    # Cria novos paragrafos como copia do anchor (pra herdar estilo Normal)
    # mas com texto novo
    for i, linha in enumerate(linhas):
        # Trata bullet markdown
        is_bullet = linha.startswith("- ") or linha.startswith("* ")
        if is_bullet:
            linha = linha[2:].strip()

        # Trata blockquote
        if linha.startswith("> "):
            linha = linha[2:].strip()

        novo_p = doc.add_paragraph()
        novo_p.text = linha
        if is_bullet:
            try:
                novo_p.style = doc.styles["List Bullet"]
            except KeyError:
                pass  # estilo nao existe, deixa como Normal

        # Move o novo paragrafo pra posicao certa (logo apos anchor + paragrafos ja inseridos)
        novo_elem = novo_p._element
        novo_elem.getparent().remove(novo_elem)
        parent.insert(indice_anchor + 1 + i, novo_elem)


def preencher_template(template_path: Path, dados: dict, output_path: Path):
    """
    Preenche o template no lugar:
    1. Copia template pra output (template original intacto).
    2. Para cada Heading 3 do template, busca conteudo no MD pelo nome.
    3. Apaga prompts (paragrafos normal abaixo do Heading 3).
    4. Insere conteudo real no lugar.
    """
    Document, qn = carregar_python_docx()

    # 1. Copia template
    shutil.copy2(template_path, output_path)
    print(f"Template copiado: {template_path.name} -> {output_path.name}")

    # 2. Abre a copia
    doc = Document(str(output_path))

    paragrafos = list(doc.paragraphs)
    total_paragrafos = len(paragrafos)

    campos_preenchidos = 0
    campos_nao_encontrados = []

    # Itera pelos paragrafos identificando Heading 1/2/3
    # Pra cada Heading 3, busca o conteudo no MD e substitui os prompts abaixo
    i = 0
    while i < len(paragrafos):
        p = paragrafos[i]
        style = p.style.name if p.style else "Normal"

        if style == "Heading 3":
            nome_campo = p.text.strip()
            nome_norm = normalizar(nome_campo)
            if not nome_norm:
                i += 1
                continue

            # Busca conteudo no MD (em qualquer h3)
            conteudo = buscar_conteudo(dados, nome_norm, nivel="h3")

            if conteudo is None or not conteudo.strip():
                campos_nao_encontrados.append(nome_campo)
                i += 1
                continue

            # Identifica paragrafos prompt a apagar (normal style entre este H3 e proximo Heading)
            j = i + 1
            prompts_a_apagar = []
            while j < len(paragrafos):
                p_seguinte = paragrafos[j]
                style_seguinte = p_seguinte.style.name if p_seguinte.style else "Normal"
                if style_seguinte in ("Heading 1", "Heading 2", "Heading 3"):
                    break
                # Apaga paragrafos com texto (que sao os prompts) e tambem vazios
                prompts_a_apagar.append(p_seguinte)
                j += 1

            # Insere conteudo real logo apos o Heading 3 (antes de apagar)
            inserir_conteudo_apos(p, conteudo, doc)

            # Agora apaga os prompts originais
            for prompt_p in prompts_a_apagar:
                try:
                    deletar_paragrafo(prompt_p)
                except Exception as e:
                    pass  # ja foi deletado ou nao existe mais

            campos_preenchidos += 1

            # Re-pega lista de paragrafos pois mudou
            paragrafos = list(doc.paragraphs)
            # Reposiciona i no proximo Heading
            try:
                novo_i = paragrafos.index(p) + 1
                i = novo_i
            except ValueError:
                i += 1
        else:
            i += 1

    # Substitui titulo da capa (primeiro Heading 1/Title)
    titulo_md = dados.get("titulo", "")
    if titulo_md:
        # Remove "Offerbook —" do inicio se houver, deixa so o nome
        nome_oferta = re.sub(r"^offerbook\s*[-–—]\s*", "", titulo_md, flags=re.IGNORECASE)
        for p in doc.paragraphs[:5]:
            style = p.style.name if p.style else ""
            if style == "Title" and p.text.strip():
                # Substitui "Nome da Oferta (ou produto)" pelo nome real
                if "Nome da Oferta" in p.text or "(ou produto)" in p.text:
                    p.text = f"Livro da Oferta (Story Selling)\n{nome_oferta}"
                    break

    # 3. Salva
    doc.save(str(output_path))
    print(f"\n✅ DOCX gerado com sucesso: {output_path}")
    print(f"   Campos preenchidos: {campos_preenchidos}")
    if campos_nao_encontrados:
        print(f"   Campos não encontrados no MD: {len(campos_nao_encontrados)}")
        for c in campos_nao_encontrados[:10]:
            print(f"     - {c}")
        if len(campos_nao_encontrados) > 10:
            print(f"     ... e mais {len(campos_nao_encontrados) - 10}")
        print("\n   (Esses campos ficaram com o prompt original do template ou em branco.")
        print("    Adicione no MD com o título exato pra preencher na próxima execução.)")


def main():
    parser = argparse.ArgumentParser(
        description="Preenche Template-Offerbook.docx no lugar a partir de um MD de offerbook."
    )
    parser.add_argument("md_file", help="Arquivo MD do offerbook (gerado pela skill /offerbook)")
    parser.add_argument(
        "--output",
        help="Caminho do DOCX de saída (default: mesmo nome do MD com .docx)",
    )
    parser.add_argument(
        "--template",
        help="Caminho customizado do Template-Offerbook.docx (default: busca automatica)",
    )
    args = parser.parse_args()

    md_path = Path(args.md_file).expanduser().resolve()
    output_path = (
        Path(args.output).expanduser().resolve()
        if args.output
        else md_path.with_suffix(".docx")
    )

    print(f"MD: {md_path}")
    print(f"Output: {output_path}\n")

    template = encontrar_template(args.template)
    print(f"Template: {template}\n")

    dados = parsear_md(md_path)
    print(f"Título: {dados['titulo']}")
    print(f"Blocos H1 encontrados: {list(dados['h1'].keys())}\n")

    preencher_template(template, dados, output_path)

    print(f"\nPronto. Abra com: open '{output_path}'")


if __name__ == "__main__":
    main()
